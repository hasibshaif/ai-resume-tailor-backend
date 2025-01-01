import os
import openai
import tempfile
import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.shared import Twips

load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")


def chunk_text(text, max_tokens=1000):
    """
    Chunk text into parts that fit within the token limit.
    """
    words = text.split()
    chunks = []
    current_chunk = []

    for word in words:
        if len(" ".join(current_chunk)) + len(word) + 1 > max_tokens:
            chunks.append(" ".join(current_chunk))
            current_chunk = []
        current_chunk.append(word)

    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks


def extract_docx_structure(file_path):
    doc = Document(file_path)
    content = []

    doc_props = {
        "headers": [
            {"text": section.header.paragraphs[0].text if section.header.paragraphs else ""}
            for section in doc.sections
        ],
        "footers": [
            {"text": section.footer.paragraphs[0].text if section.footer.paragraphs else ""}
            for section in doc.sections
        ],
        "section_props": [
            {
                "page_width": section.page_width.twips,
                "page_height": section.page_height.twips,
                "left_margin": section.left_margin.twips,
                "right_margin": section.right_margin.twips,
                "top_margin": section.top_margin.twips,
                "bottom_margin": section.bottom_margin.twips,
                "header_distance": section.header_distance.twips,
                "footer_distance": section.footer_distance.twips,
            }
            for section in doc.sections
        ],
    }

    for para in doc.paragraphs:
        para_props = {
            "text": para.text,
            "style": para.style.name,
            "alignment": para.paragraph_format.alignment,
            "runs": [
                {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None,
                }
                for run in para.runs
            ],
        }
        content.append(para_props)

    return content, doc_props


def apply_formatting(doc, content, doc_props):
    for idx, section_props in enumerate(doc_props["section_props"]):
        section = doc.sections[idx] if idx < len(doc.sections) else doc.add_section()
        section.page_width = Twips(section_props["page_width"])
        section.page_height = Twips(section_props["page_height"])
        section.left_margin = Twips(section_props["left_margin"])
        section.right_margin = Twips(section_props["right_margin"])
        section.top_margin = Twips(section_props["top_margin"])
        section.bottom_margin = Twips(section_props["bottom_margin"])
        section.header_distance = Twips(section_props["header_distance"])
        section.footer_distance = Twips(section_props["footer_distance"])

        if idx < len(doc_props["headers"]):
            header = section.header
            if doc_props["headers"][idx]["text"]:
                header.paragraphs[0].text = doc_props["headers"][idx]["text"]
        if idx < len(doc_props["footers"]):
            footer = section.footer
            if doc_props["footers"][idx]["text"]:
                footer.paragraphs[0].text = doc_props["footers"][idx]["text"]

    for para_props in content:
        if para_props["text"].startswith(("*", "-")):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(para_props["text"][1:].strip())
        else:
            paragraph = doc.add_paragraph(para_props["text"])

        paragraph.alignment = para_props["alignment"]

        for run_props in para_props["runs"]:
            run = paragraph.add_run(run_props["text"])
            run.bold = run_props["bold"]
            run.italic = run_props["italic"]
            run.underline = run_props["underline"]
            if run_props["font_name"]:
                run.font.name = run_props["font_name"]
            if run_props["font_size"]:
                run.font.size = Pt(run_props["font_size"])


def generate_tailored_resume_with_chunking(
    s3_url, job_title, job_description, output_file_path
):
    response = requests.get(s3_url)
    if response.status_code != 200:
        raise Exception("Failed to download master resume.")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
        temp_file.write(response.content)
        local_path = temp_file.name

    content, doc_props = extract_docx_structure(local_path)
    master_text = "\n".join([p["text"] for p in content])
    chunks = chunk_text(master_text)

    tailored_content = []
    for chunk in chunks:
        prompt = f"""
        Rewrite this resume content to match the job title and description below while maintaining the original formatting:
        - Ensure headers and footers are retained.
        - Preserve bullet points as actual bullets.
        - Match the formatting of the input text.

        Resume Content:
        {chunk}

        Job Title: {job_title}

        Job Description: {job_description}

        Provide only the tailored resume content, formatted exactly like the input.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0.1,
        )
        tailored_content.append(response["choices"][0]["message"]["content"])

    tailored_doc = Document()
    apply_formatting(tailored_doc, content, doc_props)

    tailored_doc.save(output_file_path)